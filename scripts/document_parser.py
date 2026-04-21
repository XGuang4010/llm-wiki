#!/usr/bin/env python3
"""
document_parser.py

DOCX document parser for llm-wiki skill.
Extracts text, headings, tables, and images from .docx files,
converts to markdown with metadata sidecar.
"""

import argparse
import hashlib
import io
import re
import shutil
import sys
import yaml
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


# Heading style mapping: Word style name -> Markdown level
HEADING_STYLE_MAP = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
    "Title": 1,
    "Subtitle": 2,
}


@dataclass
class ExtractedImage:
    """Represents an extracted image from a docx document."""
    index: int
    blob: bytes
    ext: str  # png, jpeg, etc.
    filename: str = ""
    description: str = ""


@dataclass
class TableCell:
    """Represents a table cell with content."""
    text: str
    row_span: int = 1
    col_span: int = 1
    is_merged_start: bool = True  # Is this the first cell of a merged region


@dataclass
class ParsedTable:
    """Represents a parsed table with normalized cells."""
    rows: int
    cols: int
    cells: List[List[TableCell]] = field(default_factory=list)
    has_merged_cells: bool = False


@dataclass
class DocumentMeta:
    """Metadata for a parsed document."""
    title: str = ""
    author: str = ""
    subject: str = ""
    created: str = ""
    modified: str = ""
    filename: str = ""
    source_type: str = "docx"
    word_count: int = 0
    paragraph_count: int = 0
    table_count: int = 0
    image_count: int = 0
    tags: List[str] = field(default_factory=list)
    custom: Dict[str, Any] = field(default_factory=dict)
    # New fields for data integrity and element indexing
    checksum: str = ""  # sha256:...
    size_bytes: int = 0
    elements: Dict[str, List[Dict[str, Any]]] = field(default_factory=dict)


def get_style_name(paragraph: Paragraph) -> str:
    """Get the style name of a paragraph."""
    try:
        if paragraph.style and paragraph.style.name:
            return paragraph.style.name
    except Exception:
        pass
    return "Normal"


def is_heading(paragraph: Paragraph) -> Tuple[bool, int]:
    """
    Check if paragraph is a heading and return (is_heading, level).
    Level is 1-6 for h1-h6, 0 if not a heading.
    """
    style_name = get_style_name(paragraph)

    # Check exact match
    if style_name in HEADING_STYLE_MAP:
        return True, HEADING_STYLE_MAP[style_name]

    # Check for localized heading names (e.g., "标题 1" in Chinese)
    heading_pattern = r"^(Heading|标题)\s*(\d+)"
    match = re.match(heading_pattern, style_name, re.IGNORECASE)
    if match:
        level = int(match.group(2))
        if 1 <= level <= 6:
            return True, level

    # Check if starts with Title/Subtitle patterns
    if style_name.lower().startswith("title"):
        return True, 1
    if style_name.lower().startswith("subtitle"):
        return True, 2

    return False, 0


def paragraph_to_markdown(paragraph: Paragraph) -> str:
    """Convert a paragraph to markdown text."""
    text = paragraph.text.strip()
    if not text:
        return ""

    # Check for heading
    is_h, level = is_heading(paragraph)
    if is_h:
        return f"{'#' * level} {text}"

    # Check for list items
    # Check if paragraph has list formatting
    try:
        if paragraph._p.pPr is not None:
            numPr = paragraph._p.pPr.numPr
            if numPr is not None:
                # This is a numbered or bulleted list item
                # For simplicity, we use - for bullets and 1. for numbered
                ilvl = 0
                try:
                    ilvl_elem = numPr.ilvl
                    if ilvl_elem is not None:
                        ilvl = int(ilvl_elem.val)
                except Exception:
                    pass
                indent = "  " * ilvl
                return f"{indent}- {text}"
    except Exception:
        pass

    # Regular paragraph
    return text


def extract_images(doc: DocumentType, docx_path: Path) -> List[ExtractedImage]:
    """Extract all images from a docx document."""
    images: List[ExtractedImage] = []

    # Collect all relationships for images
    rels = doc.part.rels

    image_index = 0
    for rel in rels.values():
        # Check if this relationship is an image
        if "image" in rel.target_ref:
            try:
                image_part = rel.target_part
                blob = image_part.blob

                # Determine image type from content type
                content_type = getattr(image_part, 'content_type', '')
                if 'png' in content_type or blob[:8] == b'\x89PNG\r\n\x1a\n':
                    ext = 'png'
                elif 'jpeg' in content_type or blob[:3] == b'\xff\xd8\xff':
                    ext = 'jpg'
                elif 'gif' in content_type or blob[:6] in [b'GIF87a', b'GIF89a']:
                    ext = 'gif'
                elif 'bmp' in content_type or blob[:2] == b'BM':
                    ext = 'bmp'
                else:
                    ext = 'png'  # Default to png

                img = ExtractedImage(
                    index=image_index,
                    blob=blob,
                    ext=ext,
                    filename=f"img_{image_index:02d}.{ext}"
                )
                images.append(img)
                image_index += 1
            except Exception as e:
                # Skip problematic images
                continue

    return images


def parse_table(table: Table) -> ParsedTable:
    """
    Parse a docx table, handling merged cells.
    Merged cells are split into individual cells with content in the first cell.
    """
    rows = len(table.rows)
    cols = len(table.columns)

    parsed = ParsedTable(rows=rows, cols=cols)
    parsed.cells = [[TableCell("") for _ in range(cols)] for _ in range(rows)]

    # Track merged cells
    # We'll create a grid and mark merged cells
    grid: List[List[Optional[Tuple[int, int]]]] = [[None for _ in range(cols)] for _ in range(rows)]

    for row_idx, row in enumerate(table.rows):
        col_idx = 0
        seen_tcs = set()  # Track already-processed cell elements in this row
        for cell in row.cells:
            # Find next available column
            while col_idx < cols and grid[row_idx][col_idx] is not None:
                col_idx += 1

            if col_idx >= cols:
                break

            # Skip if we've already processed this cell element (merged cell)
            tc_id = id(cell._tc)
            if tc_id in seen_tcs:
                continue
            seen_tcs.add(tc_id)

            # Get cell properties
            tc = cell._tc
            tcPr = tc.tcPr

            row_span = 1
            col_span = 1

            if tcPr is not None:
                # Check for vertical merge
                vMerge = tcPr.vMerge
                if vMerge is not None:
                    val = getattr(vMerge, 'val', None)
                    if val is None:
                        # Continue merge - handled by grid tracking
                        row_span = 1
                    else:
                        # Restart merge
                        # Need to count how many rows
                        pass

                # Check for gridSpan (horizontal merge)
                gridSpan = tcPr.gridSpan
                if gridSpan is not None:
                    val = getattr(gridSpan, 'val', None)
                    if val is not None:
                        col_span = int(val)

            # Get actual merge values from XML
            try:
                # Parse the XML to get proper merge info
                from docx.oxml import parse_xml
                tc_xml = parse_xml(tc.xml)

                # Check vMerge
                v_merge_elem = tc_xml.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
                if v_merge_elem is not None:
                    v_merge_val = v_merge_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if v_merge_val is None:
                        # This cell is part of a vertical merge but not the start
                        row_span = 0

                # Check gridSpan
                grid_span_elem = tc_xml.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
                if grid_span_elem is not None:
                    span_val = grid_span_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if span_val:
                        col_span = int(span_val)

            except Exception:
                pass

            # Store cell content
            cell_text = cell.text.strip()
            is_start = True

            # Check if this cell continues a vertical merge from above
            if row_idx > 0 and grid[row_idx][col_idx] is not None:
                # This position is already occupied by a merge
                # Find the actual row/col that started this merge
                origin_row, origin_col = grid[row_idx][col_idx]
                if origin_row == row_idx - 1 and origin_col == col_idx:
                    # Continue vertical merge
                    is_start = False
                    row_span = 0  # Mark as merged

            # Mark grid positions
            for r in range(row_idx, min(row_idx + (1 if row_span == 0 else row_span), rows)):
                for c in range(col_idx, min(col_idx + col_span, cols)):
                    if r < rows and c < cols:
                        grid[r][c] = (row_idx, col_idx)

            parsed.cells[row_idx][col_idx] = TableCell(
                text=cell_text,
                row_span=max(row_span, 1),
                col_span=col_span,
                is_merged_start=is_start
            )

            # Mark has_merged_cells
            if col_span > 1 or (row_span > 1 or not is_start):
                parsed.has_merged_cells = True

            col_idx += col_span

    # Second pass: handle vertical merge content propagation
    for row_idx in range(rows):
        for col_idx in range(cols):
            cell = parsed.cells[row_idx][col_idx]
            if not cell.is_merged_start and row_idx > 0:
                # Find the cell above that started this merge
                for r in range(row_idx - 1, -1, -1):
                    above_cell = parsed.cells[r][col_idx]
                    if above_cell.is_merged_start:
                        # Copy content to the starting cell (accumulate if different)
                        if cell.text and cell.text not in above_cell.text:
                            if above_cell.text:
                                above_cell.text += " " + cell.text
                            else:
                                above_cell.text = cell.text
                        parsed.cells[row_idx][col_idx] = TableCell(
                            text="",
                            row_span=1,
                            col_span=1,
                            is_merged_start=False
                        )
                        break

    return parsed


def table_to_markdown(table: Table, table_index: int) -> str:
    """Convert a docx table to markdown format."""
    parsed = parse_table(table)
    lines: List[str] = []

    for row_idx, row in enumerate(parsed.cells):
        row_cells: List[str] = []
        for col_idx, cell in enumerate(row):
            # Escape pipe characters in cell text
            text = cell.text.replace('|', '\\|').replace('\n', '<br>')
            row_cells.append(text)

        lines.append('| ' + ' | '.join(row_cells) + ' |')

        # Add header separator after first row
        if row_idx == 0:
            separators = ['---'] * len(row)
            lines.append('| ' + ' | '.join(separators) + ' |')

    # Add note about merged cells
    if parsed.has_merged_cells:
        lines.append("")
        lines.append("> 注：原表格有合并单元格，已按规则拆分。合并区域文字已合并到左上角第一个单元格。")

    return '\n'.join(lines)


def extract_document_meta(doc: DocumentType, docx_path: Path) -> DocumentMeta:
    """Extract metadata from a docx document."""
    meta = DocumentMeta()

    # Core properties
    try:
        core_props = doc.core_properties
        meta.title = core_props.title or ""
        meta.author = core_props.author or ""
        meta.subject = core_props.subject or ""

        if core_props.created:
            meta.created = core_props.created.isoformat()
        if core_props.modified:
            meta.modified = core_props.modified.isoformat()
    except Exception:
        pass

    meta.filename = docx_path.name

    # Count elements
    meta.paragraph_count = len(doc.paragraphs)
    meta.table_count = len(doc.tables)

    # Count words (approximate)
    word_count = 0
    for para in doc.paragraphs:
        word_count += len(para.text.split())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                word_count += len(cell.text.split())
    meta.word_count = word_count

    return meta


def parse_docx(docx_path: Path, slug: Optional[str] = None) -> Tuple[str, DocumentMeta, List[ExtractedImage]]:
    """
    Parse a docx file and return (markdown_content, metadata, images).
    Also records element positions (tables and images) with heading context.
    """
    doc = Document(docx_path)

    # Extract metadata
    meta = extract_document_meta(doc, docx_path)

    # Extract images
    images = extract_images(doc, docx_path)
    meta.image_count = len(images)

    # Build markdown content
    md_lines: List[str] = []
    table_index = 0

    # Track if we need to add title from first heading
    title_from_heading = ""

    # We need to interleave paragraphs and tables in document order
    body_elements = doc.element.body
    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)

    current_heading = ""
    table_elements: List[Dict[str, Any]] = []
    image_contexts_scanned: List[str] = []

    # Better approach: iterate through body children
    for element in body_elements:
        tag = element.tag.split('}')[-1]  # Remove namespace

        if tag == 'p':
            # Find matching paragraph
            for para in paragraphs:
                if para._p is element:
                    # Update heading context
                    is_h, level = is_heading(para)
                    if is_h:
                        current_heading = para.text.strip()
                        if level == 1 and not title_from_heading and not meta.title:
                            title_from_heading = current_heading

                    # Check for inline images in this paragraph
                    para_has_image = False
                    for run in para.runs:
                        if run._r is not None:
                            drawings = run._r.findall(qn('w:drawing'))
                            if drawings:
                                para_has_image = True
                                break

                    if para_has_image:
                        para_text = para.text.strip()
                        preview = para_text[:50] if para_text else ""
                        if current_heading:
                            if preview:
                                ctx = f"{current_heading}, {preview}..."
                            else:
                                ctx = f"{current_heading}, 段落之后"
                        else:
                            ctx = preview if preview else "文档内嵌图片"
                        image_contexts_scanned.append(ctx)

                    md_text = paragraph_to_markdown(para)
                    if md_text:
                        md_lines.append(md_text)
                    break

        elif tag == 'tbl':
            # Find matching table
            for table in tables:
                if table._tbl is element:
                    # Build table context
                    if current_heading:
                        if md_lines and not md_lines[-1].startswith('#'):
                            ctx = f"{current_heading}, 段落之后"
                        else:
                            ctx = f"{current_heading}, 段落之前"
                    else:
                        ctx = "表格"

                    md_table = table_to_markdown(table, table_index)
                    # Record table element info
                    first_line = md_table.split('\n')[0]
                    table_elements.append({
                        "index": table_index,
                        "original_range": f"table:{table_index}",
                        "markdown_ref": first_line[:80],
                        "context": ctx,
                    })

                    if md_lines and md_lines[-1] != "":
                        md_lines.append("")
                    md_lines.append(md_table)
                    md_lines.append("")
                    table_index += 1
                    break

    # Set title if not in metadata
    if not meta.title and title_from_heading:
        meta.title = title_from_heading

    meta.table_count = table_index

    # Build image element index
    image_elements: List[Dict[str, Any]] = []
    for idx, img in enumerate(images):
        if idx < len(image_contexts_scanned):
            ctx = image_contexts_scanned[idx]
        else:
            ctx = current_heading if current_heading else "文档图片"
        image_elements.append({
            "index": idx,
            "file": f"raw/images/{slug}_img_{idx:02d}.{img.ext}",
            "context": ctx,
        })

    meta.elements = {
        "tables": table_elements,
        "images": image_elements,
    }

    content = '\n\n'.join(md_lines)
    return content, meta, images


def save_images(images: List[ExtractedImage], images_dir: Path, slug: str) -> List[str]:
    """
    Save extracted images to disk and return list of relative paths.
    """
    saved_paths: List[str] = []

    for img in images:
        # Rename with slug prefix
        filename = f"{slug}_img_{img.index:02d}.{img.ext}"
        img_path = images_dir / filename

        with open(img_path, 'wb') as f:
            f.write(img.blob)

        saved_paths.append(f"raw/images/{filename}")

    return saved_paths


def meta_to_yaml(meta: DocumentMeta, slug: str) -> str:
    """Convert DocumentMeta to YAML string in the structured format."""
    data = {
        "title": meta.title or slug,
        "slug": slug,
        "type": "document",
        "source_type": meta.source_type,
        "filename": meta.filename,
        "source": {
            "path": f"raw/documents/{slug}.docx",
            "checksum": meta.checksum if meta.checksum else "",
            "size_bytes": meta.size_bytes,
        },
        "conversion": {
            "timestamp": meta.created or datetime.now().isoformat(),
            "tool": "document_parser.py",
        },
        "stats": {
            "word_count": meta.word_count,
            "paragraph_count": meta.paragraph_count,
            "table_count": meta.table_count,
            "image_count": meta.image_count,
        },
        "elements": meta.elements if meta.elements else {"tables": [], "images": []},
        "author": meta.author,
        "tags": meta.tags,
    }

    if meta.custom:
        data["custom"] = meta.custom

    return yaml.dump(data, allow_unicode=True, sort_keys=False, default_flow_style=False)


def main():
    parser = argparse.ArgumentParser(description="Parse DOCX to markdown")
    parser.add_argument("input_file", type=Path, help="Path to input .docx file")
    parser.add_argument("--slug", type=str, help="Document slug (default: filename stem)")
    parser.add_argument("--output-dir", type=Path, help="Output directory")
    parser.add_argument("--images-dir", type=Path, help="Images output directory")
    parser.add_argument("--meta-only", action="store_true", help="Only output metadata")

    args = parser.parse_args()

    if not args.input_file.exists():
        print(f"Error: File not found: {args.input_file}", file=sys.stderr)
        return 1

    slug = args.slug or args.input_file.stem

    # Parse document
    content, meta, images = parse_docx(args.input_file, slug)

    # Output directory handling
    output_dir = args.output_dir or Path(".")
    output_dir.mkdir(parents=True, exist_ok=True)

    images_dir = args.images_dir or output_dir / "raw" / "images"
    images_dir.mkdir(parents=True, exist_ok=True)

    # Save images
    if images:
        image_paths = save_images(images, images_dir, slug)
        # Add image references to content
        for i, path in enumerate(image_paths):
            # Insert image reference at appropriate position (simplified: append)
            content += f"\n\n![Image {i}]({path})"

    # Output markdown
    md_path = output_dir / f"{slug}.md"
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"Markdown saved: {md_path}")

    # Output metadata
    meta_path = output_dir / f"{slug}.meta.yaml"
    meta_yaml = meta_to_yaml(meta, slug)
    with open(meta_path, 'w', encoding='utf-8') as f:
        f.write(meta_yaml)
    print(f"Metadata saved: {meta_path}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
